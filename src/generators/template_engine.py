# 🧪 自定义模板引擎 - 支持用户上传模板
# Custom Template Engine - Support User-uploaded Templates

import os
import re
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass
import json

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class TemplateField:
    """模板字段定义"""
    name: str           # 字段名，如 {{title}}
    field_type: str     # text, table, image, chart
    required: bool = True
    description: str = ""
    placeholder: str = ""


@dataclass
class UserTemplate:
    """用户上传的模板"""
    name: str
    file_path: str
    template_type: str  # word, html, markdown
    fields: List[TemplateField] = None
    variables: Dict[str, str] = None
    
    def __post_init__(self):
        if self.fields is None:
            self.fields = []
        if self.variables is None:
            self.variables = {}


class TemplateEngine:
    """自定义模板引擎 - 解析和填充用户模板"""
    
    SUPPORTED_TYPES = ['.docx', '.html', '.md', '.markdown']
    
    def __init__(self):
        self.templates = {}
        self.current_template = None
    
    def load_template(self, file_path: str) -> UserTemplate:
        """加载用户模板"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"模板文件不存在: {file_path}")
        
        if path.suffix.lower() not in self.SUPPORTED_TYPES:
            raise ValueError(f"不支持的模板格式: {path.suffix}")
        
        # 检测模板类型
        template_type = path.suffix.lower().replace('.', '')
        if template_type in ['md', 'markdown']:
            template_type = 'markdown'
        
        # 解析字段
        fields = self._parse_fields(file_path, template_type)
        
        template = UserTemplate(
            name=path.stem,
            file_path=str(path),
            template_type=template_type,
            fields=fields,
            variables=self._parse_variables(file_path, template_type)
        )
        
        self.templates[file_path] = template
        self.current_template = template
        
        return template
    
    def _parse_fields(self, file_path: str, template_type: str) -> List[TemplateField]:
        """解析模板中的字段"""
        fields = []
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 查找 {{variable}} 格式
        pattern = r'\{\{(\w+)\}\}'
        matches = re.findall(pattern, content)
        
        for var in set(matches):
            field = TemplateField(
                name=var,
                field_type='text',
                required=True,
                description=f"变量 {var}"
            )
            fields.append(field)
        
        return fields
    
    def _parse_variables(self, file_path: str, template_type: str) -> Dict[str, str]:
        """解析模板变量"""
        variables = {}
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 查找 {{variable}} 格式
        pattern = r'\{\{(\w+)\}\}'
        matches = re.findall(pattern, content)
        
        for var in set(matches):
            variables[var] = f"[{var}]"
        
        return variables
    
    def fill_template(self, template: UserTemplate, data: Dict[str, Any]) -> str:
        """填充模板"""
        with open(template.file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 替换变量
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            content = content.replace(placeholder, str(value))
        
        return content
    
    def save_filled(self, content: str, output_path: str):
        """保存填充后的模板"""
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"✅ 模板已保存: {output_path}")
        return output_path


class WordTemplateEngine:
    """Word 模板引擎 - 支持书签和表格"""
    
    def __init__(self, template_path: str = None):
        self.doc = None
        self.bookmarks = {}
        self.tables = []
        self.template_path = template_path
    
    def load_template(self, file_path: str):
        """加载 Word 模板"""
        self.template_path = file_path
        self.doc = Document(file_path)
        self._find_bookmarks()
        self._find_tables()
        return self
    
    def _find_bookmarks(self):
        """查找 Word 书签"""
        self.bookmarks = {}
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath('.//a:bookmarkStart'):
                    for elem in run._element.xpath('.//a:bookmarkStart'):
                        name = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
                        self.bookmarks[name] = {
                            'paragraph': paragraph,
                            'text': ''
                        }
    
    def _find_tables(self):
        """查找表格"""
        self.tables = []
        for i, table in enumerate(self.doc.tables):
            self.tables.append({
                'index': i,
                'rows': len(table.rows),
                'columns': len(table.columns)
            })
    
    def fill_bookmark(self, name: str, text: str):
        """填充书签"""
        if name in self.bookmarks:
            para = self.bookmarks[name]['paragraph']
            # 清除原有内容
            para.clear()
            # 添加新内容
            run = para.add_run(text)
            self.bookmarks[name]['text'] = text
            return True
        return False
    
    def fill_text(self, pattern: str, text: str):
        """填充文本（简单替换）"""
        count = 0
        for paragraph in self.doc.paragraphs:
            if pattern in paragraph.text:
                # 替换文本
                new_text = paragraph.text.replace(pattern, text)
                paragraph.clear()
                paragraph.add_run(new_text)
                count += 1
        return count
    
    def add_table_row(self, table_index: int, data: List[str]):
        """在表格中添加行"""
        if table_index < len(self.doc.tables):
            table = self.doc.tables[table_index]
            row = table.add_row()
            for i, cell_text in enumerate(data):
                if i < len(row.cells):
                    row.cells[i].text = cell_text
            return True
        return False
    
    def insert_image(self, paragraph_index: int, image_path: str, width: Inches = Inches(5)):
        """在段落插入图片"""
        if paragraph_index < len(self.doc.paragraphs):
            para = self.doc.paragraphs[paragraph_index]
            para.add_run().add_picture(image_path, width=width)
            return True
        return False
    
    def save(self, output_path: str):
        """保存文档"""
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        print(f"✅ Word 文档已保存: {output_path}")
        return output_path
    
    def get_bookmarks(self) -> List[str]:
        """获取所有书签"""
        return list(self.bookmarks.keys())


# 便捷函数
def fill_user_template(template_path: str, data: Dict, output: str) -> str:
    """填充用户模板"""
    engine = TemplateEngine()
    template = engine.load_template(template_path)
    content = engine.fill_template(template, data)
    engine.save_filled(content, output)
    return output


if __name__ == "__main__":
    # 测试
    data = {
        "title": "欧姆定律验证实验",
        "author": "张三",
        "group": "物理1班第3组",
        "date": "2026-02-09",
        "conclusion": "实验结果符合欧姆定律预期"
    }
    
    # 创建一个测试模板
    test_template = """
# {{title}} 实验报告

**作者**: {{author}}
**组别**: {{group}}
**日期**: {{date}}

## 实验结论

{{conclusion}}
"""
    
    Path("test_template.md").write_text(test_template)
    
    # 填充
    output = fill_user_template("test_template.md", data, "test_filled.md")
    print(f"✅ 已生成: {output}")
    
    # 清理
    Path("test_template.md").unlink()
    Path("test_filled.md").unlink()
