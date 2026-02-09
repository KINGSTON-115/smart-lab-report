# 🧪 Word 报告生成器 - 支持.docx格式
# Word Report Generator - Support .docx format

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Any, Optional
from pathlib import Path
import base64
from io import BytesIO
import pandas as pd

from .chart_generator import ChartGenerator, ChartConfig

class WordReportGenerator:
    """Word 报告生成器 - 生成 .docx 格式实验报告"""
    
    def __init__(self, template_name: str = "physics_basic"):
        self.template_name = template_name
        self.doc = Document()
        self.charts = []
        self._setup_fonts()
    
    def _setup_fonts(self):
        """设置默认字体"""
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(12)
    
    def _add_heading(self, text: str, level: int = 1):
        """添加标题"""
        heading = self.doc.add_heading('', level)
        run = heading.add_run(text)
        run.font.size = Pt(14 + (2 - level) * 2)
        run.font.bold = True
    
    def _add_paragraph(self, text: str = "", style: str = None):
        """添加段落"""
        para = self.doc.add_paragraph(text, style=style)
        return para
    
    def _add_page_break(self):
        """添加分页"""
        self.doc.add_page_break()
    
    def _add_image_from_base64(self, image_base64: str, width: Inches = Inches(6)):
        """从 base64 添加图片"""
        # 解码 base64
        header, encoded = image_base64.split(',', 1)
        image_data = base64.b64decode(encoded)
        
        # 保存到临时文件
        with open('temp_chart.png', 'wb') as f:
            f.write(image_data)
        
        # 添加图片
        self.doc.add_picture('temp_chart.png', width=width)
        
        # 删除临时文件
        import os
        os.remove('temp_chart.png')
    
    def _save_image(self, image_base64: str, path: str):
        """保存图片到文件"""
        header, encoded = image_base64.split(',', 1)
        image_data = base64.b64decode(encoded)
        with open(path, 'wb') as f:
            f.write(image_data)
    
    def generate_report(self, title: str, author: str = "", group: str = "",
                       date: str = "", conclusion: str = "",
                       data_summary: Dict = None, charts: List[Dict] = None):
        """生成完整实验报告"""
        
        # 标题
        self._add_heading(title, level=0)
        
        # 元信息
        meta_para = self._add_paragraph()
        if author:
            meta_para.add_run(f"作者: {author}    ")
        if group:
            meta_para.add_run(f"组别: {group}    ")
        if date:
            meta_para.add_run(f"日期: {date}")
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 根据模板生成内容
        self._generate_content_by_template(conclusion, data_summary, charts)
        
        return self.doc
    
    def _generate_content_by_template(self, conclusion: str, 
                                      data_summary: Dict, charts: List[Dict]):
        """根据模板生成内容"""
        
        if self.template_name == "physics_basic":
            self._generate_physics_content(conclusion, data_summary, charts)
        elif self.template_name == "chemistry_basic":
            self._generate_chemistry_content(conclusion, data_summary, charts)
        elif self.template_name == "biology_basic":
            self._generate_biology_content(conclusion, data_summary, charts)
        elif self.template_name == "cs_algorithm":
            self._generate_cs_content(conclusion, data_summary, charts)
        else:
            self._generate_default_content(conclusion, data_summary, charts)
    
    def _generate_physics_content(self, conclusion: str, data_summary: Dict, charts: List[Dict]):
        """生成物理实验内容"""
        sections = [
            ("一、实验目的", "请填写实验目的..."),
            ("二、实验原理", "请填写实验原理..."),
            ("三、实验仪器", "请填写实验仪器..."),
            ("四、实验步骤", "请填写实验步骤..."),
            ("五、数据处理", self._render_data_table(data_summary)),
            ("六、误差分析", "请分析误差来源..."),
            ("七、结论与讨论", conclusion or "请填写结论..."),
        ]
        
        for title, content in sections:
            self._add_heading(title, level=1)
            if title.startswith("五、数据处理") and charts:
                # 添加图表
                for chart in charts:
                    if 'image_base64' in chart:
                        self._add_image_from_base64(chart['image_base64'])
                        self._add_paragraph()
            self._add_paragraph(content)
    
    def _generate_chemistry_content(self, conclusion: str, data_summary: Dict, charts: List[Dict]):
        """生成化学实验内容"""
        sections = [
            ("一、实验目的", "请填写实验目的..."),
            ("二、实验原理", "请填写实验原理..."),
            ("三、试剂与仪器", "请填写试剂与仪器..."),
            ("四、实验步骤", "请填写实验步骤..."),
            ("五、数据与观察", self._render_data_table(data_summary)),
            ("六、计算", "请填写计算过程..."),
            ("七、误差分析", "请分析误差来源..."),
            ("八、结论", conclusion or "请填写结论..."),
        ]
        
        for title, content in sections:
            self._add_heading(title, level=1)
            if "数据" in title and charts:
                for chart in charts:
                    if 'image_base64' in chart:
                        self._add_image_from_base64(chart['image_base64'])
            self._add_paragraph(content)
    
    def _generate_biology_content(self, conclusion: str, data_summary: Dict, charts: List[Dict]):
        """生成生物实验内容"""
        sections = [
            ("一、实验目的", "请填写实验目的..."),
            ("二、背景介绍", "请填写背景..."),
            ("三、材料与方法", "请填写材料与方法..."),
            ("四、实验结果", self._render_data_table(data_summary)),
            ("五、分析讨论", conclusion or "请填写分析..."),
            ("六、结论", "请填写结论..."),
        ]
        
        for title, content in sections:
            self._add_heading(title, level=1)
            self._add_paragraph(content)
    
    def _generate_cs_content(self, conclusion: str, data_summary: Dict, charts: List[Dict]):
        """生成计算机实验内容"""
        sections = [
            ("一、问题描述", "请描述问题..."),
            ("二、算法设计", "请填写算法设计..."),
            ("三、时间复杂度分析", "请分析复杂度..."),
            ("四、实现代码", "```python\n# 请粘贴代码\n```"),
            ("五、测试用例", self._render_data_table(data_summary)),
            ("六、实验结果", "请填写实验结果..."),
            ("七、讨论与优化", conclusion or "请填写讨论..."),
        ]
        
        for title, content in sections:
            self._add_heading(title, level=1)
            self._add_paragraph(content)
    
    def _generate_default_content(self, conclusion: str, data_summary: Dict, charts: List[Dict]):
        """默认内容"""
        self._add_heading("一、实验数据", level=1)
        self._add_paragraph(self._render_data_table(data_summary))
        
        if charts:
            for chart in charts:
                if 'image_base64' in chart:
                    self._add_image_from_base64(chart['image_base64'])
        
        self._add_heading("二、实验结论", level=1)
        self._add_paragraph(conclusion or "请填写结论...")
    
    def _render_data_table(self, data_summary: Dict) -> str:
        """渲染数据表格"""
        if not data_summary or not data_summary.get("statistics"):
            return "暂无数据"
        
        # 添加表格
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '列名'
        header_cells[1].text = '均值'
        header_cells[2].text = '标准差'
        header_cells[3].text = '变异系数(%)'
        
        # 数据行
        for col, stats in data_summary["statistics"].items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(col)
            row_cells[1].text = f"{stats.get('mean', 0):.4f}"
            row_cells[2].text = f"{stats.get('std', 0):.4f}"
            row_cells[3].text = f"{stats.get('cv', 0):.2f}%"
        
        return "数据表格已生成"
    
    def save(self, output_path: str):
        """保存 Word 文档"""
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        print(f"✅ Word 报告已保存: {output_path}")
        return output_path


# 便捷函数
def generate_word_report(data_path: str, title: str, author: str = "", 
                         group: str = "", template: str = "physics_basic",
                         output: str = "output/report.docx") -> str:
    """快速生成 Word 报告"""
    from src.generators.report_generator import ReportGenerator
    
    # 加载数据
    data = pd.read_csv(data_path) if data_path.endswith('.csv') else pd.read_excel(data_path)
    
    # 生成图表
    chart_gen = ChartGenerator(data)
    numeric_cols = data.select_dtypes(include=['number']).columns.tolist()
    
    charts = []
    if len(numeric_cols) >= 2:
        result = chart_gen.generate(numeric_cols[0], [numeric_cols[1]], 
                                   ChartConfig(title=f"{numeric_cols[0]} vs {numeric_cols[1]}"))
        charts.append(result)
    
    # 生成报告
    word_gen = WordReportGenerator(template)
    doc = word_gen.generate_report(
        title=title,
        author=author,
        group=group,
        conclusion="请根据实验结果填写结论...",
        data_summary={"statistics": {}}
    )
    
    # 保存
    word_gen.save(output)
    return output


if __name__ == "__main__":
    # 测试
    output = generate_word_report(
        data_path="data/examples/欧姆定律数据.csv",
        title="欧姆定律验证实验",
        author="张三",
        group="物理1班第3组",
        output="output/欧姆定律实验报告.docx"
    )
    print(f"✅ Word 报告生成成功: {output}")
